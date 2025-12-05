#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
数据获取与管理模块

支持多种文件格式的处理，包括PDF、Word、Excel、CSV、网页抓取等。
提供数据验证、质量检查、批处理和存储管理功能。

Author: KG-CLIR Team
"""

import os
import io
import json
import sqlite3
import hashlib
import pandas as pd
import requests
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlparse, urljoin
import tempfile
import shutil

# 文件处理库
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup
import openpyxl

# 项目内部模块
from src.utils.io import load_jsonl, save_jsonl
from src.utils.text_norm import normalize_text
from src.utils.lang_detect import detect_language
from src.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class DocumentMetadata:
    """文档元数据"""
    doc_id: str
    filename: str
    file_type: str
    size: int
    language: str
    created_at: datetime
    processed_at: datetime
    quality_score: float
    checksum: str
    source: str = "upload"

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "doc_id": self.doc_id,
            "filename": self.filename,
            "file_type": self.file_type,
            "size": self.size,
            "language": self.language,
            "created_at": self.created_at.isoformat(),
            "processed_at": self.processed_at.isoformat(),
            "quality_score": self.quality_score,
            "checksum": self.checksum,
            "source": self.source
        }


class FileProcessor:
    """文件处理器"""

    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # 简化处理
            '.txt': self._process_text,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.json': self._process_json,
            '.jsonl': self._process_jsonl
        }

    def _calculate_checksum(self, file_path: Union[str, Path]) -> str:
        """计算文件校验和"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _process_pdf(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理PDF文件"""
        try:
            # 首先尝试使用pdfplumber (更好的表格和布局支持)
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
            except Exception:
                # 降级到PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"

            metadata = {
                "pages": len(PyPDF2.PdfReader(open(file_path, 'rb')).pages) if os.path.exists(file_path) else 0,
                "processor": "pdfplumber" if "pdfplumber" in str(type(pdfplumber)) else "PyPDF2"
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"PDF处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def _process_docx(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理Word文档"""
        try:
            doc = Document(file_path)
            text = ""

            # 处理段落
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # 处理表格
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    tables_text.append(row_text)
                tables_text.append("")  # 表格间空行

            if tables_text:
                text += "\n".join(tables_text)

            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "processor": "python-docx"
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"Word文档处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def _process_text(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理纯文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # 尝试其他编码
            if not text:
                encodings = ['gbk', 'gb2312', 'latin-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue

            metadata = {
                "lines": len(text.splitlines()),
                "characters": len(text),
                "processor": "text"
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"文本文件处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def _process_csv(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理CSV文件"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            df = None

            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                raise ValueError("无法读取CSV文件，尝试了多种编码")

            # 转换为文本
            text = df.to_string(index=False)

            metadata = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "processor": "pandas"
            }

            return text, metadata

        except Exception as e:
            logger.error(f"CSV文件处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def _process_excel(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理Excel文件"""
        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            sheet_info = {}

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_text = df.to_string(index=False)
                text_parts.append(f"=== 工作表: {sheet_name} ===\n{sheet_text}")

                sheet_info[sheet_name] = {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist()
                }

            text = "\n\n".join(text_parts)

            metadata = {
                "sheets": len(excel_file.sheet_names),
                "sheet_names": excel_file.sheet_names,
                "sheet_info": sheet_info,
                "processor": "openpyxl"
            }

            return text, metadata

        except Exception as e:
            logger.error(f"Excel文件处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def _process_json(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # 提取所有文本内容
            def extract_text(obj):
                if isinstance(obj, str):
                    return obj
                elif isinstance(obj, dict):
                    return " ".join(extract_text(v) for v in obj.values())
                elif isinstance(obj, list):
                    return " ".join(extract_text(item) for item in obj)
                else:
                    return str(obj)

            text = extract_text(data)

            metadata = {
                "type": type(data).__name__,
                "processor": "json"
            }

            if isinstance(data, dict):
                metadata["keys"] = list(data.keys())
            elif isinstance(data, list):
                metadata["length"] = len(data)

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"JSON文件处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def _process_jsonl(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """处理JSONL文件"""
        try:
            data = load_jsonl(file_path)

            # 提取所有文本内容
            text_parts = []
            for item in data:
                if isinstance(item, dict):
                    # 寻找常见的文本字段
                    text_fields = ['text', 'content', 'description', 'title', 'abstract']
                    item_texts = []
                    for field in text_fields:
                        if field in item and isinstance(item[field], str):
                            item_texts.append(item[field])

                    if item_texts:
                        text_parts.append(" ".join(item_texts))
                    else:
                        # 如果没有常见字段，提取所有字符串值
                        all_texts = []
                        for v in item.values():
                            if isinstance(v, str):
                                all_texts.append(v)
                        if all_texts:
                            text_parts.append(" ".join(all_texts))

            text = "\n\n".join(text_parts)

            metadata = {
                "records": len(data),
                "processor": "jsonl"
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"JSONL文件处理失败: {file_path}, 错误: {e}")
            return "", {"error": str(e)}

    def process_file(self, file_path: Union[str, Path]) -> Tuple[Optional[str], Optional[Dict[str, Any]], Optional[str]]:
        """
        处理单个文件

        Returns:
            Tuple[文本内容, 元数据, 错误信息]
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return None, None, f"文件不存在: {file_path}"

        # 检查文件格式
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_formats:
            return None, None, f"不支持的文件格式: {file_ext}"

        try:
            # 调用对应的处理器
            processor = self.supported_formats[file_ext]
            text, metadata = processor(file_path)

            if not text.strip():
                return None, None, "文件内容为空或处理失败"

            # 基础元数据
            base_metadata = {
                "filename": file_path.name,
                "file_size": file_path.stat().st_size,
                "file_type": file_ext,
                "checksum": self._calculate_checksum(file_path)
            }
            base_metadata.update(metadata)

            return text, base_metadata, None

        except Exception as e:
            logger.error(f"文件处理异常: {file_path}, 错误: {e}")
            return None, None, f"处理异常: {str(e)}"


class DataValidator:
    """数据验证器"""

    def __init__(self):
        self.min_text_length = 50
        self.max_text_length = 1000000  # 1M字符
        self.supported_languages = ['zh', 'fr', 'en', 'unknown']

    def validate_text(self, text: str) -> Dict[str, Any]:
        """验证文本内容"""
        validation_result = {
            "is_valid": True,
            "issues": [],
            "quality_score": 1.0
        }

        # 长度检查
        text_length = len(text.strip())
        if text_length < self.min_text_length:
            validation_result["is_valid"] = False
            validation_result["issues"].append(f"文本过短: {text_length} 字符")
            validation_result["quality_score"] *= 0.3

        if text_length > self.max_text_length:
            validation_result["issues"].append(f"文本过长: {text_length} 字符")
            validation_result["quality_score"] *= 0.8

        # 语言检查
        try:
            language = detect_language(text)
            if language not in self.supported_languages:
                validation_result["issues"].append(f"不支持的语言: {language}")
                validation_result["quality_score"] *= 0.7
            validation_result["language"] = language
        except Exception as e:
            validation_result["language"] = "unknown"
            validation_result["issues"].append(f"语言检测失败: {e}")
            validation_result["quality_score"] *= 0.9

        # 内容质量检查
        # 检查是否主要是空白字符
        non_whitespace_ratio = len(text.strip()) / max(len(text), 1)
        if non_whitespace_ratio < 0.5:
            validation_result["issues"].append("空白字符比例过高")
            validation_result["quality_score"] *= 0.6

        # 检查重复内容
        lines = text.splitlines()
        if len(lines) > 10:
            unique_lines = len(set(line.strip() for line in lines if line.strip()))
            duplicate_ratio = 1 - (unique_lines / len(lines))
            if duplicate_ratio > 0.7:
                validation_result["issues"].append(f"重复内容过多: {duplicate_ratio:.2%}")
                validation_result["quality_score"] *= 0.7

        return validation_result

    def check_duplicates(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """检查重复文档"""
        duplicates = {}
        seen_checksums = {}

        for i, doc in enumerate(documents):
            checksum = doc.get('checksum', '')
            if checksum:
                if checksum in seen_checksums:
                    if checksum not in duplicates:
                        duplicates[checksum] = [seen_checksums[checksum]]
                    duplicates[checksum].append(i)
                else:
                    seen_checksums[checksum] = i

        return duplicates


class BatchProcessor:
    """批处理器"""

    def __init__(self, max_workers: int = 4):
        self.max_workers = max_workers
        self.file_processor = FileProcessor()
        self.data_validator = DataValidator()

    def process_files(self, file_paths: List[Union[str, Path]],
                     progress_callback=None) -> List[Dict[str, Any]]:
        """批量处理文件"""
        results = []

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # 提交所有任务
            future_to_file = {
                executor.submit(self._process_single_file, file_path): file_path
                for file_path in file_paths
            }

            # 收集结果
            for i, future in enumerate(as_completed(future_to_file)):
                file_path = future_to_file[future]
                try:
                    result = future.result(timeout=60)  # 60秒超时
                    results.append(result)

                    if progress_callback:
                        progress_callback(i + 1, len(file_paths), file_path.name)

                except Exception as e:
                    logger.error(f"批处理失败: {file_path}, 错误: {e}")
                    results.append({
                        "filename": Path(file_path).name,
                        "status": "error",
                        "error": str(e)
                    })

        return results

    def _process_single_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """处理单个文件（内部方法）"""
        file_path = Path(file_path)

        # 处理文件
        text, metadata, error = self.file_processor.process_file(file_path)

        if error:
            return {
                "filename": file_path.name,
                "status": "error",
                "error": error
            }

        # 验证数据
        validation = self.data_validator.validate_text(text)

        # 生成文档ID
        doc_id = hashlib.md5(f"{file_path.name}_{metadata.get('checksum', '')}".encode()).hexdigest()[:16]

        # 创建文档元数据
        doc_metadata = DocumentMetadata(
            doc_id=doc_id,
            filename=file_path.name,
            file_type=metadata.get("file_type", ""),
            size=metadata.get("file_size", 0),
            language=validation.get("language", "unknown"),
            created_at=datetime.fromtimestamp(file_path.stat().st_ctime),
            processed_at=datetime.now(),
            quality_score=validation["quality_score"],
            checksum=metadata.get("checksum", ""),
            source="upload"
        )

        # 预处理文本
        processed_text = normalize_text(text)

        return {
            "doc_id": doc_id,
            "filename": file_path.name,
            "status": "success",
            "text": processed_text,
            "original_text": text,
            "metadata": doc_metadata.to_dict(),
            "processing_metadata": metadata,
            "validation": validation
        }


class DataStorage:
    """数据存储管理"""

    def __init__(self, storage_path: Union[str, Path] = "data/uploads"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)

        # SQLite数据库路径
        self.db_path = self.storage_path / "documents.db"
        self._init_database()

    def _init_database(self):
        """初始化数据库"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    doc_id TEXT PRIMARY KEY,
                    filename TEXT,
                    file_type TEXT,
                    size INTEGER,
                    language TEXT,
                    created_at TEXT,
                    processed_at TEXT,
                    quality_score REAL,
                    checksum TEXT,
                    source TEXT,
                    metadata TEXT,
                    processing_metadata TEXT,
                    validation TEXT
                )
            """)

            conn.execute("""
                CREATE TABLE IF NOT EXISTS document_content (
                    doc_id TEXT PRIMARY KEY,
                    text TEXT,
                    original_text TEXT,
                    FOREIGN KEY (doc_id) REFERENCES documents (doc_id)
                )
            """)

            conn.commit()

    def save_document(self, doc_data: Dict[str, Any]) -> bool:
        """保存文档到数据库"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # 保存文档元数据
                metadata = json.dumps(doc_data.get("metadata", {}))
                processing_metadata = json.dumps(doc_data.get("processing_metadata", {}))
                validation = json.dumps(doc_data.get("validation", {}))

                conn.execute("""
                    INSERT OR REPLACE INTO documents
                    (doc_id, filename, file_type, size, language, created_at,
                     processed_at, quality_score, checksum, source, metadata,
                     processing_metadata, validation)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    doc_data["doc_id"],
                    doc_data["filename"],
                    doc_data["metadata"]["file_type"],
                    doc_data["metadata"]["size"],
                    doc_data["metadata"]["language"],
                    doc_data["metadata"]["created_at"],
                    doc_data["metadata"]["processed_at"],
                    doc_data["metadata"]["quality_score"],
                    doc_data["metadata"]["checksum"],
                    doc_data["metadata"]["source"],
                    metadata,
                    processing_metadata,
                    validation
                ))

                # 保存文档内容
                conn.execute("""
                    INSERT OR REPLACE INTO document_content (doc_id, text, original_text)
                    VALUES (?, ?, ?)
                """, (
                    doc_data["doc_id"],
                    doc_data["text"],
                    doc_data.get("original_text", "")
                ))

                conn.commit()
                return True

        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            return False

    def get_documents(self, limit: int = 100, offset: int = 0,
                     language: str = None, min_quality: float = None) -> List[Dict[str, Any]]:
        """获取文档列表"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                query = """
                    SELECT d.*, dc.text, dc.original_text
                    FROM documents d
                    LEFT JOIN document_content dc ON d.doc_id = dc.doc_id
                    WHERE 1=1
                """
                params = []

                if language:
                    query += " AND d.language = ?"
                    params.append(language)

                if min_quality:
                    query += " AND d.quality_score >= ?"
                    params.append(min_quality)

                query += " ORDER BY d.processed_at DESC LIMIT ? OFFSET ?"
                params.extend([limit, offset])

                cursor = conn.execute(query, params)
                rows = cursor.fetchall()

                columns = [desc[0] for desc in cursor.description]
                documents = []

                for row in rows:
                    doc = dict(zip(columns, row))
                    # 解析JSON字段
                    for field in ['metadata', 'processing_metadata', 'validation']:
                        if doc[field]:
                            try:
                                doc[field] = json.loads(doc[field])
                            except:
                                pass
                    documents.append(doc)

                return documents

        except Exception as e:
            logger.error(f"获取文档列表失败: {e}")
            return []

    def get_document_by_id(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """根据ID获取文档"""
        documents = self.get_documents(limit=1)
        for doc in documents:
            if doc['doc_id'] == doc_id:
                return doc
        return None

    def delete_document(self, doc_id: str) -> bool:
        """删除文档"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute("DELETE FROM document_content WHERE doc_id = ?", (doc_id,))
                conn.execute("DELETE FROM documents WHERE doc_id = ?", (doc_id,))
                conn.commit()
                return True
        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            return False

    def get_statistics(self) -> Dict[str, Any]:
        """获取统计信息"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # 总文档数
                total_docs = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]

                # 语言分布
                language_stats = conn.execute("""
                    SELECT language, COUNT(*) as count
                    FROM documents
                    GROUP BY language
                """).fetchall()

                # 文件类型分布
                file_type_stats = conn.execute("""
                    SELECT file_type, COUNT(*) as count
                    FROM documents
                    GROUP BY file_type
                """).fetchall()

                # 平均质量分数
                avg_quality = conn.execute("SELECT AVG(quality_score) FROM documents").fetchone()[0]

                return {
                    "total_documents": total_docs,
                    "language_distribution": dict(language_stats),
                    "file_type_distribution": dict(file_type_stats),
                    "average_quality_score": avg_quality or 0.0
                }

        except Exception as e:
            logger.error(f"获取统计信息失败: {e}")
            return {}


class WebScraper:
    """网页抓取器"""

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })

    def scrape_url(self, url: str, timeout: int = 30) -> Tuple[str, Dict[str, Any]]:
        """抓取单个URL"""
        try:
            response = self.session.get(url, timeout=timeout)
            response.raise_for_status()

            # 解析HTML
            soup = BeautifulSoup(response.content, 'html.parser')

            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()

            # 提取文本
            title = soup.title.string if soup.title else ""
            text = soup.get_text(separator='\n', strip=True)

            # 清理文本
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)

            metadata = {
                "url": url,
                "title": title,
                "status_code": response.status_code,
                "content_type": response.headers.get('content-type', ''),
                "content_length": len(text),
                "processor": "beautifulsoup"
            }

            return text, metadata

        except Exception as e:
            logger.error(f"网页抓取失败: {url}, 错误: {e}")
            return "", {"error": str(e)}

    def scrape_urls(self, urls: List[str], progress_callback=None) -> List[Dict[str, Any]]:
        """批量抓取URLs"""
        results = []

        for i, url in enumerate(urls):
            try:
                text, metadata = self.scrape_url(url)

                if text.strip():
                    # 语言检测
                    language = detect_language(text[:1000])  # 使用前1000字符检测语言

                    # 创建文档数据
                    doc_id = hashlib.md5(f"url_{url}".encode()).hexdigest()[:16]

                    doc_data = {
                        "doc_id": doc_id,
                        "filename": f"web_{doc_id}.txt",
                        "status": "success",
                        "text": normalize_text(text),
                        "original_text": text,
                        "source": "web_scraping",
                        "language": language,
                        "url": url,
                        "metadata": metadata
                    }

                    results.append(doc_data)
                else:
                    results.append({
                        "url": url,
                        "status": "error",
                        "error": "无内容或抓取失败"
                    })

                if progress_callback:
                    progress_callback(i + 1, len(urls), url)

            except Exception as e:
                logger.error(f"URL处理异常: {url}, 错误: {e}")
                results.append({
                    "url": url,
                    "status": "error",
                    "error": str(e)
                })

        return results


# 便利函数
def create_data_manager(storage_path: str = "data/uploads") -> DataStorage:
    """创建数据管理器实例"""
    return DataStorage(storage_path)


def process_uploaded_files(files, storage_path: str = "data/uploads",
                          progress_callback=None) -> List[Dict[str, Any]]:
    """处理上传的文件（Streamlit专用）"""
    batch_processor = BatchProcessor()
    data_storage = DataStorage(storage_path)

    # 保存上传的文件到临时目录
    temp_dir = Path(storage_path) / "temp"
    temp_dir.mkdir(exist_ok=True)

    file_paths = []
    for file in files:
        # 保存文件
        file_path = temp_dir / file.name
        with open(file_path, "wb") as f:
            f.write(file.getbuffer())
        file_paths.append(file_path)

    try:
        # 批处理文件
        results = batch_processor.process_files(file_paths, progress_callback)

        # 保存到数据库
        saved_results = []
        for result in results:
            if result.get("status") == "success":
                if data_storage.save_document(result):
                    saved_results.append(result)
                else:
                    result["status"] = "error"
                    result["error"] = "保存到数据库失败"
                    saved_results.append(result)
            else:
                saved_results.append(result)

        return saved_results

    finally:
        # 清理临时文件
        for file_path in file_paths:
            try:
                file_path.unlink()
            except:
                pass


if __name__ == "__main__":
    # 简单测试
    storage = create_data_manager()
    stats = storage.get_statistics()
    print("数据存储统计:", stats)