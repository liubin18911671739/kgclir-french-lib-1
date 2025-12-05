"""
文本处理器模块

提供多语言文本处理和元数据提取功能，包括：
- 文本预处理和规范化
- 多语言支持和语言检测
- 元数据提取和结构化
- 知识提取和实体识别
"""

import re
import json
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Dict, List, Optional, Any, Tuple, Union
from pathlib import Path
import hashlib

# 外部库
import spacy
import jieba
import langdetect
from bs4 import BeautifulSoup
import pdfplumber
import PyPDF2
from docx import Document
import pandas as pd

# 本地导入
from ..utils.text_norm import TextNormalizer
from ..utils.lang_detect import LanguageDetector

# 设置日志
logger = logging.getLogger(__name__)

@dataclass
class ProcessedText:
    """处理后的文本数据"""
    original_text: str
    normalized_text: str
    language: str
    confidence: float
    metadata: Dict[str, Any]
    entities: List[Dict[str, Any]]
    sentences: List[str]
    words: List[str]
    word_count: int
    sentence_count: int
    quality_score: float

@dataclass
class ExtractedMetadata:
    """提取的元数据"""
    title: Optional[str] = None
    authors: List[str] = None
    publication_date: Optional[str] = None
    abstract: Optional[str] = None
    keywords: List[str] = None
    doi: Optional[str] = None
    url: Optional[str] = None
    language: Optional[str] = None
    word_count: int = 0
    page_count: int = 0
    references: List[str] = None
    citations: List[str] = None
    tables: List[Dict[str, Any]] = None
    figures: List[Dict[str, Any]] = None

    def __post_init__(self):
        if self.authors is None:
            self.authors = []
        if self.keywords is None:
            self.keywords = []
        if self.references is None:
            self.references = []
        if self.citations is None:
            self.citations = []
        if self.tables is None:
            self.tables = []
        if self.figures is None:
            self.figures = []

class BaseProcessor(ABC):
    """处理器基类"""

    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.text_normalizer = TextNormalizer(config.get('text_normalization', {}))
        self.language_detector = LanguageDetector(config.get('language_detection', {}))

    @abstractmethod
    def process(self, text: str, **kwargs) -> ProcessedText:
        """处理文本"""
        pass

    @abstractmethod
    def extract_metadata(self, text: str, **kwargs) -> ExtractedMetadata:
        """提取元数据"""
        pass

class TextProcessor(BaseProcessor):
    """多语言文本处理器"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.nlp_models = {}
        self._load_nlp_models()

    def _load_nlp_models(self):
        """加载NLP模型"""
        model_configs = self.config.get('nlp_models', {
            'zh': {'model': 'zh_core_web_sm', 'enabled': True},
            'fr': {'model': 'fr_core_news_sm', 'enabled': True},
            'en': {'model': 'en_core_web_sm', 'enabled': True}
        })

        for lang, config in model_configs.items():
            if config.get('enabled', True):
                try:
                    model_name = config['model']
                    self.nlp_models[lang] = spacy.load(model_name)
                    logger.info(f"加载 {lang} NLP模型: {model_name}")
                except OSError:
                    logger.warning(f"未找到 {lang} NLP模型: {model_name}")
                except Exception as e:
                    logger.error(f"加载 {lang} NLP模型失败: {e}")

    def process(self, text: str, language: str = None, **kwargs) -> ProcessedText:
        """处理文本"""
        if not text or not text.strip():
            return self._create_empty_processed_text()

        # 检测语言
        if language is None:
            detected_lang, confidence = self.language_detector.detect(text)
        else:
            detected_lang, confidence = language, 1.0

        # 文本规范化
        normalized_text = self.text_normalizer.normalize(text)

        # 分词和句子分割
        sentences = self._extract_sentences(normalized_text, detected_lang)
        words = self._extract_words(normalized_text, detected_lang)

        # 实体识别
        entities = self._extract_entities(normalized_text, detected_lang)

        # 质量评分
        quality_score = self._calculate_quality_score(
            text, normalized_text, entities, sentences, words
        )

        # 生成元数据
        metadata = {
            'processing_language': detected_lang,
            'language_confidence': confidence,
            'original_length': len(text),
            'normalized_length': len(normalized_text),
            'compression_ratio': len(normalized_text) / len(text) if text else 0,
            'has_entities': len(entities) > 0,
            'processing_timestamp': self._get_timestamp()
        }

        return ProcessedText(
            original_text=text,
            normalized_text=normalized_text,
            language=detected_lang,
            confidence=confidence,
            metadata=metadata,
            entities=entities,
            sentences=sentences,
            words=words,
            word_count=len(words),
            sentence_count=len(sentences),
            quality_score=quality_score
        )

    def extract_metadata(self, text: str, **kwargs) -> ExtractedMetadata:
        """提取文本元数据"""
        metadata = ExtractedMetadata()

        # 提取标题
        metadata.title = self._extract_title(text)

        # 提取作者
        metadata.authors = self._extract_authors(text)

        # 提取关键词
        metadata.keywords = self._extract_keywords(text)

        # 提取摘要
        metadata.abstract = self._extract_abstract(text)

        # 提取参考文献
        metadata.references = self._extract_references(text)

        # 计算字数
        metadata.word_count = len(text.split())

        # 检测语言
        detected_lang, _ = self.language_detector.detect(text)
        metadata.language = detected_lang

        return metadata

    def _create_empty_processed_text(self) -> ProcessedText:
        """创建空的处理结果"""
        return ProcessedText(
            original_text="",
            normalized_text="",
            language="unknown",
            confidence=0.0,
            metadata={},
            entities=[],
            sentences=[],
            words=[],
            word_count=0,
            sentence_count=0,
            quality_score=0.0
        )

    def _extract_sentences(self, text: str, language: str) -> List[str]:
        """提取句子"""
        sentences = []

        try:
            if language in self.nlp_models:
                # 使用spaCy
                doc = self.nlp_models[language](text)
                sentences = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
            else:
                # 使用简单的正则表达式
                sentence_endings = r'[.!?。！？]+'
                raw_sentences = re.split(sentence_endings, text)
                sentences = [s.strip() for s in raw_sentences if s.strip()]
        except Exception as e:
            logger.error(f"句子提取失败: {e}")
            # 使用备用方法
            sentences = text.split('\n')
            sentences = [s.strip() for s in sentences if s.strip()]

        return sentences

    def _extract_words(self, text: str, language: str) -> List[str]:
        """提取词汇"""
        words = []

        try:
            if language == 'zh':
                # 中文使用jieba分词
                words = list(jieba.cut(text))
                # 过滤空白字符和单字符
                words = [w.strip() for w in words if w.strip() and len(w.strip()) > 1]
            elif language in self.nlp_models:
                # 使用spaCy
                doc = self.nlp_models[language](text)
                words = [token.text.strip() for token in doc if not token.is_space and not token.is_punct]
            else:
                # 使用简单的分词
                words = re.findall(r'\b\w+\b', text)
        except Exception as e:
            logger.error(f"词汇提取失败: {e}")
            # 使用备用方法
            words = text.split()
            words = [w.strip() for w in words if w.strip()]

        return words

    def _extract_entities(self, text: str, language: str) -> List[Dict[str, Any]]:
        """提取实体"""
        entities = []

        try:
            if language in self.nlp_models:
                doc = self.nlp_models[language](text)
                for ent in doc.ents:
                    entities.append({
                        'text': ent.text,
                        'label': ent.label_,
                        'start': ent.start_char,
                        'end': ent.end_char,
                        'confidence': 1.0  # spaCy不直接提供置信度
                    })
            else:
                # 使用简单的正则表达式提取一些常见实体
                entities = self._extract_entities_regex(text)
        except Exception as e:
            logger.error(f"实体提取失败: {e}")

        return entities

    def _extract_entities_regex(self, text: str) -> List[Dict[str, Any]]:
        """使用正则表达式提取实体"""
        entities = []

        # 提取电子邮件
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        for match in re.finditer(email_pattern, text):
            entities.append({
                'text': match.group(),
                'label': 'EMAIL',
                'start': match.start(),
                'end': match.end(),
                'confidence': 0.9
            })

        # 提取URL
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        for match in re.finditer(url_pattern, text):
            entities.append({
                'text': match.group(),
                'label': 'URL',
                'start': match.start(),
                'end': match.end(),
                'confidence': 0.9
            })

        # 提取年份
        year_pattern = r'\b(19|20)\d{2}\b'
        for match in re.finditer(year_pattern, text):
            entities.append({
                'text': match.group(),
                'label': 'YEAR',
                'start': match.start(),
                'end': match.end(),
                'confidence': 0.7
            })

        return entities

    def _calculate_quality_score(self, original: str, normalized: str,
                                entities: List, sentences: List, words: List) -> float:
        """计算文本质量分数"""
        score = 0.0

        if not original:
            return score

        # 长度分数 (30%)
        length_score = min(len(normalized) / 1000, 1.0) * 0.3
        score += length_score

        # 结构分数 (25%)
        if len(sentences) > 0 and len(words) > 0:
            avg_words_per_sentence = len(words) / len(sentences)
            structure_score = min(avg_words_per_sentence / 20, 1.0) * 0.25
            score += structure_score

        # 实体分数 (25%)
        entity_score = min(len(entities) / 10, 1.0) * 0.25
        score += entity_score

        # 语言一致性分数 (20%)
        try:
            detected_lang, confidence = self.language_detector.detect(normalized)
            lang_score = confidence * 0.2
            score += lang_score
        except:
            score += 0.1  # 默认分数

        return min(score, 1.0)

    def _extract_title(self, text: str) -> Optional[str]:
        """提取标题"""
        lines = text.split('\n')

        # 寻找第一行非空文本作为标题
        for line in lines[:5]:  # 只检查前5行
            line = line.strip()
            if line and len(line) > 5 and len(line) < 200:
                # 简单启发式：检查是否看起来像标题
                if not line.endswith(('.', '!', '?', '。', '！', '？')):
                    return line

        return None

    def _extract_authors(self, text: str) -> List[str]:
        """提取作者信息"""
        authors = []

        # 简单的作者模式匹配
        patterns = [
            r'(?:作者|Authors?)[:：]\s*([^。\n]+)',
            r'(?:著|by)\s+([^。\n]+)',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # 分割多个作者
                author_list = re.split(r'[,;，；]', match)
                authors.extend([a.strip() for a in author_list if a.strip()])

        return list(set(authors))  # 去重

    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        keywords = []

        # 寻找关键词部分
        patterns = [
            r'(?:关键词|Key Words?|Keywords?)[:：]\s*([^。\n]+)',
            r'(?:主题|Topics?)[:：]\s*([^。\n]+)',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # 分割关键词
                keyword_list = re.split(r'[,;，；\s]+', match)
                keywords.extend([k.strip() for k in keyword_list if k.strip() and len(k.strip()) > 1])

        return list(set(keywords))  # 去重

    def _extract_abstract(self, text: str) -> Optional[str]:
        """提取摘要"""
        # 寻找摘要部分
        patterns = [
            r'(?:摘要|Abstract)[:：]\s*([^。\n]*(?:\n[^。\n]*)*?)(?=\n\s*[A-Z]|\n\s*关键词|\Z)',
            r'(?:概要|Summary)[:：]\s*([^。\n]*(?:\n[^。\n]*)*?)(?=\n\s*[A-Z]|\n\s*关键词|\Z)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                abstract = match.group(1).strip()
                # 清理和验证摘要
                if len(abstract) > 50 and len(abstract) < 2000:
                    return abstract

        return None

    def _extract_references(self, text: str) -> List[str]:
        """提取参考文献"""
        references = []

        # 寻找参考文献部分
        ref_section_pattern = r'(?:参考文献|References?)[：:]?\s*(.*?)(?=\n\s*$|\n\s*[A-Z][^.\n]*\.\s|\z)'
        ref_match = re.search(ref_section_pattern, text, re.IGNORECASE | re.DOTALL)

        if ref_match:
            ref_text = ref_match.group(1).strip()

            # 分割参考文献条目
            ref_items = re.split(r'\n\s*\d+\.\s*|\n\s*\[\d+\]\s*', ref_text)

            for ref in ref_items:
                ref = ref.strip()
                if ref and len(ref) > 20:  # 过滤太短的条目
                    references.append(ref)

        return references

    def _get_timestamp(self) -> str:
        """获取当前时间戳"""
        from datetime import datetime
        return datetime.now().isoformat()

class MetadataExtractor(BaseProcessor):
    """专门的元数据提取器"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.extraction_rules = config.get('extraction_rules', {})

    def process(self, text: str, **kwargs) -> ProcessedText:
        """处理文本（主方法委托给TextProcessor）"""
        processor = TextProcessor(self.config)
        return processor.process(text, **kwargs)

    def extract_metadata(self, text: str, **kwargs) -> ExtractedMetadata:
        """提取详细元数据"""
        metadata = ExtractedMetadata()

        # 基本元数据提取
        processor = TextProcessor(self.config)
        basic_metadata = processor.extract_metadata(text, **kwargs)

        # 合并基本元数据
        metadata.title = basic_metadata.title
        metadata.authors = basic_metadata.authors
        metadata.keywords = basic_metadata.keywords
        metadata.abstract = basic_metadata.abstract
        metadata.references = basic_metadata.references
        metadata.language = basic_metadata.language
        metadata.word_count = basic_metadata.word_count

        # 高级元数据提取
        metadata.doi = self._extract_doi(text)
        metadata.url = self._extract_urls(text)
        metadata.publication_date = self._extract_publication_date(text)
        metadata.citations = self._extract_citations(text)
        metadata.tables = self._extract_tables(text)
        metadata.figures = self._extract_figures(text)

        return metadata

    def _extract_doi(self, text: str) -> Optional[str]:
        """提取DOI"""
        doi_pattern = r'(?:doi|DOI)[:：]?\s*(10\.\d+/[^\s]+)'
        match = re.search(doi_pattern, text, re.IGNORECASE)
        return match.group(1) if match else None

    def _extract_urls(self, text: str) -> List[str]:
        """提取URL"""
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        return list(set(re.findall(url_pattern, text)))

    def _extract_publication_date(self, text: str) -> Optional[str]:
        """提取发表日期"""
        # 日期模式匹配
        date_patterns = [
            r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?)',
            r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]

        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)

        return None

    def _extract_citations(self, text: str) -> List[str]:
        """提取引用"""
        citations = []

        # 文内引用模式
        citation_patterns = [
            r'\[([^\]]+)\]',  # [1], [Smith, 2020], [1-5]
            r'\(([^)]*(?:19|20)\d{2}[^)]*)\)',  # (Smith, 2020)
        ]

        for pattern in citation_patterns:
            matches = re.findall(pattern, text)
            citations.extend(matches)

        return list(set(citations))

    def _extract_tables(self, text: str) -> List[Dict[str, Any]]:
        """提取表格信息"""
        tables = []

        # 简单的表格识别模式
        table_markers = re.finditer(r'表\s*\d+[：:]?\s*([^\n]*)', text, re.IGNORECASE)

        for i, match in enumerate(table_markers):
            table_info = {
                'index': i,
                'title': match.group(1).strip(),
                'position': match.start()
            }
            tables.append(table_info)

        return tables

    def _extract_figures(self, text: str) -> List[Dict[str, Any]]:
        """提取图片信息"""
        figures = []

        # 图片识别模式
        figure_markers = re.finditer(r'图\s*\d+[：:]?\s*([^\n]*)', text, re.IGNORECASE)

        for i, match in enumerate(figure_markers):
            figure_info = {
                'index': i,
                'title': match.group(1).strip(),
                'position': match.start()
            }
            figures.append(figure_info)

        return figures

class MultilingualAligner(BaseProcessor):
    """多语言对齐器"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.alignment_methods = config.get('alignment_methods', ['lexical', 'embedding'])

    def process(self, text: str, **kwargs) -> ProcessedText:
        """处理文本"""
        processor = TextProcessor(self.config)
        return processor.process(text, **kwargs)

    def extract_metadata(self, text: str, **kwargs) -> ExtractedMetadata:
        """提取元数据"""
        processor = TextProcessor(self.config)
        return processor.extract_metadata(text, **kwargs)

    def align_texts(self, source_text: str, target_text: str,
                   source_lang: str, target_lang: str) -> Dict[str, Any]:
        """对齐两种语言的文本"""
        alignment_result = {
            'source_lang': source_lang,
            'target_lang': target_lang,
            'alignment_score': 0.0,
            'aligned_segments': [],
            'translation_confidence': 0.0
        }

        try:
            # 处理源文本和目标文本
            source_processed = self.process(source_text, language=source_lang)
            target_processed = self.process(target_text, language=target_lang)

            # 句子级别对齐
            aligned_sentences = self._align_sentences(
                source_processed.sentences, target_processed.sentences
            )

            # 计算对齐分数
            alignment_score = self._calculate_alignment_score(
                source_processed, target_processed, aligned_sentences
            )

            alignment_result.update({
                'alignment_score': alignment_score,
                'aligned_segments': aligned_sentences,
                'source_sentences': source_processed.sentences,
                'target_sentences': target_processed.sentences
            })

        except Exception as e:
            logger.error(f"文本对齐失败: {e}")

        return alignment_result

    def _align_sentences(self, source_sentences: List[str], target_sentences: List[str]) -> List[Dict[str, Any]]:
        """句子级别对齐"""
        alignments = []

        # 简单的基于长度的对齐方法
        source_len = len(source_sentences)
        target_len = len(target_sentences)

        # 计算对应关系
        ratio = target_len / source_len if source_len > 0 else 1

        for i, src_sent in enumerate(source_sentences):
            # 计算对应的句子索引
            target_idx = int(i * ratio)
            target_end_idx = min(int((i + 1) * ratio), target_len)

            if target_idx < target_len:
                if target_end_idx > target_idx + 1:
                    # 多对一对齐
                    target_sentence = ' '.join(target_sentences[target_idx:target_end_idx])
                else:
                    # 一对一对齐
                    target_sentence = target_sentences[target_idx]

                alignment = {
                    'source_index': i,
                    'target_index': target_idx,
                    'source_sentence': src_sent,
                    'target_sentence': target_sentence,
                    'alignment_type': '1-1' if target_end_idx == target_idx + 1 else '1-n'
                }
                alignments.append(alignment)

        return alignments

    def _calculate_alignment_score(self, source_processed: ProcessedText,
                                 target_processed: ProcessedText,
                                 alignments: List[Dict[str, Any]]) -> float:
        """计算对齐分数"""
        if not alignments:
            return 0.0

        # 基于对齐覆盖率的分数
        source_coverage = len(alignments) / len(source_processed.sentences) if source_processed.sentences else 0

        # 基于语言相似性的分数
        lang_similarity = 1.0 if source_processed.language == target_processed.language else 0.5

        # 基于实体重叠的分数
        source_entities = set([e['text'].lower() for e in source_processed.entities])
        target_entities = set([e['text'].lower() for e in target_processed.entities])

        if source_entities and target_entities:
            entity_overlap = len(source_entities & target_entities) / len(source_entities | target_entities)
        else:
            entity_overlap = 0.0

        # 加权平均
        score = 0.5 * source_coverage + 0.3 * lang_similarity + 0.2 * entity_overlap

        return min(score, 1.0)

class KnowledgeExtractor(BaseProcessor):
    """知识提取器"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.knowledge_patterns = config.get('knowledge_patterns', {})

    def process(self, text: str, **kwargs) -> ProcessedText:
        """处理文本"""
        processor = TextProcessor(self.config)
        return processor.process(text, **kwargs)

    def extract_metadata(self, text: str, **kwargs) -> ExtractedMetadata:
        """提取元数据"""
        processor = TextProcessor(self.config)
        return processor.extract_metadata(text, **kwargs)

    def extract_concepts(self, text: str, language: str = None) -> List[Dict[str, Any]]:
        """提取概念"""
        concepts = []

        processed = self.process(text, language=language)

        # 从实体中提取概念
        for entity in processed.entities:
            if entity['label'] in ['PERSON', 'ORG', 'GPE', 'PRODUCT', 'EVENT']:
                concept = {
                    'name': entity['text'],
                    'type': entity['label'],
                    'confidence': entity['confidence'],
                    'context': self._get_context(text, entity['start'], entity['end'])
                }
                concepts.append(concept)

        # 使用预定义模式提取概念
        pattern_concepts = self._extract_concepts_by_patterns(text, language)
        concepts.extend(pattern_concepts)

        return concepts

    def extract_relations(self, text: str, concepts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """提取关系"""
        relations = []

        # 简单的基于共现的关系提取
        for i, concept1 in enumerate(concepts):
            for j, concept2 in enumerate(concepts[i+1:], i+1):
                # 检查概念间的距离
                distance = abs(concept1.get('position', 0) - concept2.get('position', 0))

                if distance < 100:  # 在一定范围内认为是相关的
                    relation = {
                        'subject': concept1['name'],
                        'object': concept2['name'],
                        'relation_type': 'co_occurrence',
                        'confidence': 1.0 - (distance / 100),
                        'context': self._get_relation_context(text, concept1, concept2)
                    }
                    relations.append(relation)

        # 使用模式匹配提取特定关系
        pattern_relations = self._extract_relations_by_patterns(text, concepts)
        relations.extend(pattern_relations)

        return relations

    def _extract_concepts_by_patterns(self, text: str, language: str) -> List[Dict[str, Any]]:
        """使用模式提取概念"""
        concepts = []

        # 定义概念模式
        patterns = {
            'definition': [
                r'(\w+)\s*[是为]\s*([^。；;]+)',
                r'(\w+)\s*[是为]\s*([^。；;]+)',
                r'(\w+)\s*(?:refers to|is defined as|means)\s+([^。；;]+)',
            ],
            'classification': [
                r'(\w+)\s*[是属于]\s*(\w+)',
                r'(\w+)\s*[是属于]\s*(\w+)',
                r'(\w+)\s*(?:is a|is an|belongs to)\s+(\w+)',
            ]
        }

        for rel_type, pattern_list in patterns.items():
            for pattern in pattern_list:
                matches = re.finditer(pattern, text)
                for match in matches:
                    concept = {
                        'name': match.group(1),
                        'type': rel_type,
                        'confidence': 0.8,
                        'definition': match.group(2) if len(match.groups()) > 1 else None
                    }
                    concepts.append(concept)

        return concepts

    def _extract_relations_by_patterns(self, text: str, concepts: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """使用模式提取关系"""
        relations = []

        # 获取概念名称集合
        concept_names = {c['name'] for c in concepts}

        # 定义关系模式
        relation_patterns = [
            r'(\w+)\s*(?:支持|包括|包含|涵盖)\s*(\w+)',
            r'(\w+)\s*(?:requires|needs|depends on)\s+(\w+)',
            r'(\w+)\s*(?:leads to|causes|results in)\s+(\w+)',
        ]

        for pattern in relation_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                subj, obj = match.group(1), match.group(2)

                # 检查是否都是已知概念
                if subj in concept_names and obj in concept_names:
                    relation = {
                        'subject': subj,
                        'object': obj,
                        'relation_type': 'semantic_relation',
                        'confidence': 0.7,
                        'pattern': pattern
                    }
                    relations.append(relation)

        return relations

    def _get_context(self, text: str, start: int, end: int, window: int = 50) -> str:
        """获取实体的上下文"""
        context_start = max(0, start - window)
        context_end = min(len(text), end + window)
        return text[context_start:context_end].strip()

    def _get_relation_context(self, text: str, concept1: Dict, concept2: Dict) -> str:
        """获取关系的上下文"""
        pos1 = concept1.get('position', 0)
        pos2 = concept2.get('position', 0)

        if pos1 < pos2:
            return text[pos1:pos2 + 100].strip()
        else:
            return text[pos2:pos1 + 100].strip()

# 导出主要类
__all__ = [
    'TextProcessor',
    'MetadataExtractor',
    'MultilingualAligner',
    'KnowledgeExtractor',
    'ProcessedText',
    'ExtractedMetadata'
]